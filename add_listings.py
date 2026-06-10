from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

src = r"C:\Users\Windows\Downloads\СЖАТАЯ_КВАЛИФИКАЦИЯ.docx"
dst = r"C:\Users\Windows\Downloads\СЖАТАЯ_КВАЛИФИКАЦИЯ_С_ЛИСТИНГАМИ.docx"

doc = Document(src)

def add_separator():
    p = doc.add_paragraph()
    run = p.add_run("─" * 60)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    p.space_after = Pt(6)

def add_code(code_text):
    lines = code_text.split("\n")
    for line in lines:
        p = doc.add_paragraph()
        text = line if line.strip() else " "
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(11)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F0F0F0")
        shd.set(qn("w:val"), "clear")
        p._element.get_or_add_pPr().append(shd)

# ───────────────────────────────
# Listing 1 — Schema (Prisma)
# ───────────────────────────────
add_separator()
add_heading("Листинг 1 — Схема базы данных (Prisma)")

schema_user = """model User {
  id           String   @id @default(cuid())
  email        String   @unique
  passwordHash String?
  role         Role     @default(INITIATOR)
  isActive     Boolean  @default(true)
  createdAt    DateTime @default(now())
  updatedAt    DateTime @updatedAt

  employee             Employee?
  authoredDocuments    InternalDocument[]   @relation("DocumentAuthor")
  approvals            DocumentApproval[]   @relation("DocumentApprover")
  signatures           DigitalSignature[]
  notifications        Notification[]
  createdVersions      DocumentVersion[]
  incomingResolutions  IncomingDocument[]   @relation("ResolutionAuthor")
  createdIncomingDocuments IncomingDocument[] @relation("IncomingCreator")
}"""

schema_doc = """model InternalDocument {
  id         String          @id @default(cuid())
  title      String
  content    String?
  type       InternalDocType
  status     DocumentStatus  @default(DRAFT)
  number     String?
  fileUrl    String?
  deadline   DateTime?
  authorId   String
  author     User            @relation("DocumentAuthor", fields: [authorId], references: [id])
  workflowId String?
  workflow   WorkflowTemplate? @relation(fields: [workflowId], references: [id])

  approvals  DocumentApproval[]
  versions   DocumentVersion[]
  signatures DigitalSignature[]
  incomingResponseFor IncomingDocument[] @relation("IncomingResponse")

  createdAt DateTime @default(now())
  updatedAt DateTime @updatedAt
}"""

schema_approval = """model DocumentApproval {
  id          String          @id @default(cuid())
  documentId  String
  stageId     String?
  approverId  String
  decision    ReviewDecision?
  comment     String?
  signatureId String?
  decidedAt   DateTime?

  document  InternalDocument @relation(fields: [documentId], references: [id], onDelete: Cascade)
  stage     WorkflowStage?   @relation(fields: [stageId], references: [id])
  approver  User             @relation("DocumentApprover", fields: [approverId], references: [id])
  signature DigitalSignature? @relation("DocumentSignature", fields: [signatureId], references: [id])

  createdAt DateTime @default(now())

  @@index([documentId])
  @@index([approverId])
  @@unique([signatureId])
}

model DigitalSignature {
  id            String   @id @default(cuid())
  documentId    String
  employeeId    String
  userId        String
  signatureData String
  documentHash  String
  algorithm     String   @default("aes-256-gcm")
  isVerified    Boolean  @default(false)

  document InternalDocument @relation(fields: [documentId], references: [id])
}"""

add_code(schema_user)
add_code(schema_doc)
add_code(schema_approval)

# ───────────────────────────────
# Listing 2 — Crypto
# ───────────────────────────────
add_separator()
add_heading("Листинг 2 — Реализация AES-256-GCM шифрования")

crypto_code = """import { createCipheriv, createDecipheriv, randomBytes } from "crypto";

const ALGORITHM = "aes-256-gcm";

function getKey(): Buffer {
  const key = process.env.ENCRYPTION_KEY || "";
  if (!key) {
    throw new Error("ENCRYPTION_KEY not set in environment");
  }
  return Buffer.from(key, "hex");
}

export function encrypt(text: string): string {
  const key = getKey();
  const iv = randomBytes(16);
  
  const cipher = createCipheriv(ALGORITHM, key, iv);
  let encrypted = cipher.update(text, "utf8", "hex");
  encrypted += cipher.final("hex");
  
  const authTag = cipher.getAuthTag().toString("hex");
  
  return `${iv.toString("hex")}:${authTag}:${encrypted}`;
}

export function decrypt(encryptedText: string): string {
  const key = getKey();
  const [ivHex, authTagHex, encrypted] = encryptedText.split(":");
  
  if (!ivHex || !authTagHex || !encrypted) {
    throw new Error("Invalid encrypted text format");
  }
  
  const decipher = createDecipheriv(ALGORITHM, key, Buffer.from(ivHex, "hex"));
  decipher.setAuthTag(Buffer.from(authTagHex, "hex"));
  
  let decrypted = decipher.update(encrypted, "hex", "utf8");
  decrypted += decipher.final("utf8");
  
  return decrypted;
}"""

add_code(crypto_code)

# ───────────────────────────────
# Listing 3 — Create + SendToWorkflow
# ───────────────────────────────
add_separator()
add_heading("Листинг 3 — Создание документа и отправка на согласование")

create_doc_code = """export async function createDocument(data: {
  title: string;
  content?: string;
  type: InternalDocType;
  fileUrl?: string;
}) {
  const session = await auth();
  if (!session?.user) throw new Error("Не авторизован");

  const template = await prisma.workflowTemplate.findFirst({
    where: { docType: data.type },
    include: { stages: { orderBy: { stageOrder: "asc" } } },
  });

  const doc = await prisma.internalDocument.create({
    data: {
      title: data.title,
      content: data.content,
      type: data.type,
      status: "DRAFT",
      fileUrl: data.fileUrl,
      authorId: session.user.id,
      workflowId: template?.id,
    },
  });

  if (template) {
    for (const stage of template.stages) {
      const approvers = await prisma.employee.findMany({
        where: { positionId: stage.approverPositionId },
        include: { user: true },
      });
      for (const emp of approvers) {
        await prisma.documentApproval.create({
          data: {
            documentId: doc.id,
            stageId: stage.id,
            approverId: emp.user.id,
          },
        });
      }
    }
  }

  await prisma.auditLog.create({
    data: {
      userId: session.user.id,
      action: "CREATE",
      entityType: "InternalDocument",
      entityId: doc.id,
    },
  });

  revalidatePath("/documents");
  return doc;
}"""

send_to_workflow_code = """export async function sendToWorkflow(id: string) {
  const session = await auth();
  if (!session?.user) throw new Error("Не авторизован");

  const doc = await prisma.internalDocument.findUnique({
    where: { id },
    include: {
      approvals: {
        where: { stage: { stageOrder: 1 } },
        include: { approver: true },
      },
    },
  });
  if (!doc) throw new Error("Документ не найден");
  if (doc.status !== "DRAFT") throw new Error("Документ уже отправлен");

  const number = `${doc.type}-${new Date().getFullYear()}-${String(Math.floor(Math.random() * 999)).padStart(3, "0")}`;

  await prisma.internalDocument.update({
    where: { id },
    data: { status: "IN_APPROVAL", number },
  });

  for (const approval of doc.approvals) {
    await createNotification(
      approval.approverId,
      "APPROVAL_REQUEST",
      "Новый документ на согласование",
      `Документ "${doc.title}" ожидает вашего согласования.`,
      "InternalDocument",
      id
    );
  }

  await prisma.auditLog.create({
    data: {
      userId: session.user.id,
      action: "APPROVE",
      entityType: "InternalDocument",
      entityId: id,
      newStatus: "IN_APPROVAL",
    },
  });

  revalidatePath(`/documents/${id}`);
  return { number };
}"""

add_code(create_doc_code)
p = doc.add_paragraph()
add_code(send_to_workflow_code)

# ───────────────────────────────
# Listing 4 — Approval + Signature
# ───────────────────────────────
add_separator()
add_heading("Листинг 4 — Согласование и электронная подпись")

approval_code = """export async function submitApproval(
  approvalId: string,
  decision: "APPROVE" | "REJECT" | "RETURN_TO_AUTHOR",
  comment?: string
) {
  const session = await auth();
  if (!session?.user) throw new Error("Не авторизован");

  const approval = await prisma.documentApproval.findUnique({
    where: { id: approvalId },
    include: {
      document: {
        include: {
          workflow: {
            include: { stages: { orderBy: { stageOrder: "asc" } } },
          },
          approvals: {
            include: { approver: true },
          },
        },
      },
      stage: true,
    },
  });
  if (!approval) throw new Error("Запись согласования не найдена");
  if (approval.approverId !== session.user.id) throw new Error("Нет доступа");

  await prisma.documentApproval.update({
    where: { id: approvalId },
    data: { decision, comment, decidedAt: new Date() },
  });

  const doc = approval.document;

  if (decision === "APPROVE") {
    const currentStageOrder = approval.stage?.stageOrder || 0;
    const nextStage = doc.workflow?.stages.find((s) => s.stageOrder === currentStageOrder + 1);

    if (nextStage) {
      const nextApprovals = await prisma.documentApproval.findMany({
        where: { documentId: doc.id, stage: { stageOrder: nextStage.stageOrder } },
        include: { approver: true },
      });
      for (const na of nextApprovals) {
        await createNotification(na.approverId, "APPROVAL_REQUEST",
          "Новый документ на согласование",
          `Документ "${doc.title}" ожидает вашего согласования.`,
          "InternalDocument", doc.id);
      }
    } else {
      await prisma.internalDocument.update({ where: { id: doc.id }, data: { status: "APPROVED" } });
    }
  } else if (decision === "REJECT") {
    await prisma.internalDocument.update({ where: { id: doc.id }, data: { status: "REJECTED" } });
    await createNotification(doc.authorId, "DOCUMENT_REJECTED",
      "Документ отклонён",
      `Ваш документ "${doc.title}" был отклонён.${comment ? ` Причина: ${comment}` : ""}`,
      "InternalDocument", doc.id);
  } else if (decision === "RETURN_TO_AUTHOR") {
    await prisma.internalDocument.update({ where: { id: doc.id }, data: { status: "DRAFT" } });
    await createNotification(doc.authorId, "DOCUMENT_RETURNED",
      "Документ возвращён на доработку",
      `Ваш документ "${doc.title}" возвращён на доработку.${comment ? ` Причина: ${comment}` : ""}`,
      "InternalDocument", doc.id);
  }

  await prisma.auditLog.create({
    data: { userId: session.user.id,
      action: decision === "APPROVE" ? "APPROVE" : decision === "REJECT" ? "REJECT" : "RETURN",
      entityType: "DocumentApproval", entityId: approvalId, comment,
    },
  });

  revalidatePath(`/documents/${doc.id}`);
}"""

signature_code = """export async function submitSignature(approvalId: string) {
  const session = await auth();
  if (!session?.user) throw new Error("Не авторизован");

  const approval = await prisma.documentApproval.findUnique({
    where: { id: approvalId },
    include: { document: true },
  });
  if (!approval) throw new Error("Запись согласования не найдена");
  if (approval.approverId !== session.user.id) throw new Error("Нет доступа");

  const employee = await prisma.employee.findUnique({ where: { userId: session.user.id } });
  if (!employee) throw new Error("Профиль сотрудника не найден");

  const docHash = await computeDocumentHash(approval.document.content || "");
  const encryptedSig = await signHash(docHash);

  const signature = await prisma.digitalSignature.create({
    data: {
      documentId: approval.document.id,
      employeeId: employee.id,
      userId: session.user.id,
      signatureData: encryptedSig,
      documentHash: docHash,
      isVerified: true,
    },
  });

  await prisma.documentApproval.update({
    where: { id: approvalId },
    data: { signatureId: signature.id, decidedAt: new Date(), decision: "APPROVE" },
  });

  const allApprovals = await prisma.documentApproval.findMany({
    where: { documentId: approval.document.id },
  });

  const allDecided = allApprovals.every((a) => a.decision !== null);
  if (allDecided) {
    await prisma.internalDocument.update({
      where: { id: approval.document.id },
      data: { status: "APPROVED" },
    });
  }

  await prisma.auditLog.create({
    data: { userId: session.user.id, action: "SIGN", entityType: "DigitalSignature", entityId: signature.id },
  });

  revalidatePath(`/documents/${approval.document.id}`);
  return signature;
}"""

add_code(approval_code)
p = doc.add_paragraph()
add_code(signature_code)

# ───────────────────────────────
# Listing 5 — Middleware
# ───────────────────────────────
add_separator()
add_heading("Листинг 5 — Middleware защита по ролям")

middleware_code = """import { auth } from "@/auth";
import { NextResponse } from "next/server";

const roleBasedRoutes: Record<string, string[]> = {
  "/documents/create": ["INITIATOR", "VALIDATOR", "ADMIN"],
  "/documents/pending": ["VALIDATOR", "SIGNER", "REGISTRAR", "ADMIN"],
  "/documents/approval": ["VALIDATOR", "SIGNER"],
  "/incoming/register": ["REGISTRAR", "ADMIN"],
  "/incoming": ["REGISTRAR", "ADMIN", "SIGNER", "VALIDATOR"],
  "/admin": ["ADMIN"],
};

const protectedRoutes = Object.keys(roleBasedRoutes);

export default auth((req) => {
  const path = req.nextUrl.pathname;
  const user = req.auth?.user;

  if (!user && path !== "/login" && path !== "/") {
    return NextResponse.redirect(new URL("/login", req.url));
  }

  if (user) {
    const matchedRoute = protectedRoutes.find((route) => path.startsWith(route));
    if (matchedRoute) {
      const allowedRoles = roleBasedRoutes[matchedRoute];
      if (!allowedRoles.includes(user.role as string)) {
        return NextResponse.redirect(new URL("/", req.url));
      }
    }
  }

  return NextResponse.next();
});

export const config = {
  matcher: ["/((?!api|_next/static|_next/image|images|icon|favicon|manifest).*)"],
};"""

add_code(middleware_code)

# ───────────────────────────────
# Listing 6 — ApprovalTimeline
# ───────────────────────────────
add_separator()
add_heading("Листинг 6 — Визуализация маршрута согласования")

timeline_code = """import { CheckCircle, Clock, XCircle } from "lucide-react";

interface ApprovalTimelineProps {
  workflow: any;
  approvals: any[];
  signatures: any[];
}

export function ApprovalTimeline({ workflow, approvals, signatures }: ApprovalTimelineProps) {
  const stages = workflow?.stages?.sort((a: any, b: any) => a.stageOrder - b.stageOrder) || [];

  return (
    <div className="space-y-4">
      {stages.length === 0 ? (
        <p className="text-sm text-[var(--text-muted)]">Маршрут согласования не настроен</p>
      ) : (
        stages.map((stage: any) => {
          const stageApprovals = approvals.filter((a: any) => a.stageId === stage.id);
          const decided = stageApprovals.filter((a: any) => a.decision !== null);
          const allDecided = stageApprovals.length > 0 && stageApprovals.every((a: any) => a.decision !== null);
          const isCurrent = stageApprovals.some((a: any) => a.decision === null);
          const isRejected = decided.some((a: any) => a.decision === "REJECT" || a.decision === "RETURN_TO_AUTHOR");

          return (
            <div key={stage.id} className="flex gap-4">
              <div className="flex flex-col items-center">
                <div className={`w-8 h-8 rounded-full flex items-center justify-center ${
                  allDecided && !isRejected ? "bg-[var(--success)]" :
                  isRejected ? "bg-[var(--danger)]" :
                  isCurrent ? "bg-[var(--warning)]" : "bg-[var(--bg-secondary)]"
                }`}>
                  {allDecided && !isRejected ? (
                    <CheckCircle className="w-5 h-5 text-white" />
                  ) : isRejected ? (
                    <XCircle className="w-5 h-5 text-white" />
                  ) : (
                    <span className="text-sm font-medium text-[var(--text-muted)]">{stage.stageOrder}</span>
                  )}
                </div>
                <div className="w-0.5 flex-1 bg-[var(--border-subtle)] min-h-[24px]" />
              </div>
              <div className="flex-1 pb-4">
                <p className="font-medium text-[var(--text-primary)]">{stage.approverPosition?.name || `Этап ${stage.stageOrder}`}</p>
                <p className="text-sm text-[var(--text-muted)]">
                  {stage.deadlineDays ? `Срок: ${stage.deadlineDays} дн.` : ""}
                </p>
                {stageApprovals.map((approval: any) => (
                  <div key={approval.id} className="mt-2 p-3 rounded-xl bg-[var(--bg-secondary)]">
                    <div className="flex items-center justify-between">
                      <p className="text-sm font-medium text-[var(--text-primary)]">
                        {approval.approver.employee
                          ? `${approval.approver.employee.lastName} ${approval.approver.employee.firstName}`
                          : approval.approver.email}
                      </p>
                      {approval.decision && (
                        <span className={`px-2 py-0.5 text-xs rounded-full ${
                          approval.decision === "APPROVE" ? "badge-success" :
                          approval.decision === "REJECT" ? "badge-danger" : "badge-warning"
                        }`}>
                          {approval.decision === "APPROVE" ? "Согласовано" :
                           approval.decision === "REJECT" ? "Отклонено" : "Возвращено"}
                        </span>
                      )}
                    </div>
                    {approval.comment && (
                      <p className="text-sm text-[var(--text-muted)] mt-1">{approval.comment}</p>
                    )}
                    {approval.decidedAt && (
                      <p className="text-xs text-[var(--text-muted)] mt-1">
                        {new Date(approval.decidedAt).toLocaleDateString("ru-RU")}
                      </p>
                    )}
                    {approval.signature && (
                      <p className="text-xs text-[var(--success)] mt-1">✓ Подписано ЭП</p>
                    )}
                  </div>
                ))}
                {stageApprovals.length === 0 && (
                  <p className="text-sm text-[var(--text-muted)] mt-2">Нет согласующих</p>
                )}
              </div>
            </div>
          );
        })
      )}
    </div>
  );
}"""

add_code(timeline_code)

# ───────────────────────────────
# Listing 7 — SignatureStamp
# ───────────────────────────────
add_separator()
add_heading("Листинг 7 — Штамп электронной подписи")

stamp_code = """interface SignatureStampProps {
  signature: {
    id: string;
    employee: { firstName: string; lastName: string; position?: { name: string } } | null;
    createdAt: string | Date;
    signatureData: string;
  };
}

export function SignatureStamp({ signature }: SignatureStampProps) {
  const sigShort = signature.signatureData.split(":")[0]?.slice(0, 4) || "??";
  const positionName = (signature as any).employee?.position?.name || "Сотрудник";

  return (
    <div className="border-2 border-[var(--accent)]/30 rounded-xl p-4 min-w-[200px] bg-[var(--bg-card)] shadow-lg">
      <p className="text-sm font-bold text-[var(--accent)] mb-2">УТВЕРЖДАЮ</p>
      <p className="font-medium text-[var(--text-primary)]">
        {signature.employee?.lastName} {signature.employee?.firstName}
      </p>
      <p className="text-sm text-[var(--text-muted)]">{positionName}</p>
      <p className="text-sm text-[var(--text-muted)]">{new Date(signature.createdAt).toLocaleDateString("ru-RU")}</p>
      <p className="text-xs text-[var(--accent)] mt-2 font-mono">ЭП: {sigShort}:...:12</p>
    </div>
  );
}"""

add_code(stamp_code)

doc.save(dst)
print(f"Saved to {dst}")
